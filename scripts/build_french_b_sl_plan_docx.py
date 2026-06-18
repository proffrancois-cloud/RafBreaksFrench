from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUT = Path("output") / "Plan_de_cours_Francais_B_SL_2025-2027.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(w))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = bold
    return p


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_info_table(doc):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2200, 7160])
    rows = [
        ("Document", "Plan de cours sur deux ans - Francais B SL"),
        ("Periode", "DP1 2025-2026; DP2 2026-2027, planifie d'aout a avril"),
        ("Source de travail", "Toddle: DP1/DP2 French B Unit planning et Class flow consultes le 17 juin 2026"),
        ("Intention", "Clarifier ce qui a ete couvert en DP1 et organiser DP2 de facon transversale, sans isoler artificiellement un theme IB par unite."),
    ]
    for row, (k, v) in zip(table.rows, rows):
        add_cell_text(row.cells[0], k, bold=True)
        add_cell_text(row.cells[1], v)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
    set_repeat_table_header(table.rows[0])


def add_dp1_summary(doc):
    add_heading(doc, "1. Bilan DP1 2025-2026", 1)
    add_para(
        doc,
        "La progression DP1 visible dans Toddle a commence en suivant l'aperçu transmis par l'enseignante precedente. Cette base semblait diviser les cinq themes du programme en blocs successifs, avec environ deux sous-parties par theme. Cette organisation donne une impression de couverture lineaire, mais elle n'est pas la plus efficace pour Francais B SL: les competences d'examen, les types de textes, l'oral et les ATL doivent etre revisites de facon transversale dans chaque unite.",
    )
    add_para(
        doc,
        "Pour la suite, chaque unite doit donc annoncer explicitement les themes IB et les recommended topics travailles, sans transformer chaque unite en un seul theme isole. Sinon, le temps manque et les eleves ne transferent pas assez les contenus d'un contexte a l'autre.",
    )

    rows = [
        (
            "Unit 1: Identites - Sortir de sa zone de confort\n19 aout-4 novembre; 24 h",
            "Identites: health and wellbeing; language and identity. Introduction au DP, auto-positionnement CEFR, objectifs personnels, bien-etre, Francophonie et langue comme identite.",
            "Description d'image; courriel; blog/poster/speech; opinion, sentiments, but, obligation; subjonctif, pronoms, connecteurs.",
            "Paper 1: courriel/blog. Oral: presentation du projet personnel et description visuelle.",
        ),
        (
            "Unit 2: Experiences - 2.3/2.4\n18 novembre-9 janvier; 24 h",
            "Experiences: life stories, rites of passage, migration. Travail sur les parcours personnels et culturels, lectures sur rites de passage et ressources audio/visuelles sur migration.",
            "Comprehension ecrite, comprehension orale, prefixes/suffixes, narration, cause/consequence, reformulation.",
            "Paper 2 Reading: rites de passage/migration. Paper 2 Listening: migrations et accents francophones.",
        ),
        (
            "Unit 3: Ingeniosite humaine - IA / speaking skill\n22 janvier-13 fevrier; 10 h",
            "Human ingenuity: entertainment; communication and media. Unite transversale pour preparer l'IA: arts, medias, document visuel, discussion de suivi et discussion generale.",
            "Organisation de la prise de parole, description-analyse-lien au theme, connecteurs oraux, ecoute active, strategies de preparation.",
            "Oral IA: entrainement sur support visuel. Paper 2 Listening: documents audio varies.",
        ),
        (
            "Unit 4: Organisation sociale & Partage de la planete\n13 mars-2 avril; 12 h",
            "Social organization et Sharing the planet: chomage, enjeux sociaux, mondialisation, environnement urbain/rural. Sequence courte de consolidation receptive.",
            "Paper 2 Listening et Reading, documents authentiques, pronoms relatifs, reperage des pieges et des conventions.",
            "Paper 2 Reading: livret/question booklet. Paper 2 Listening: entrainement type examen.",
        ),
        (
            "Unit 5: Guerre et paix - Ethique - Traditions\n2 avril-19 juin; 20 h",
            "Experiences, Human ingenuity, Sharing the planet: memoire de la guerre d'Algerie, traditions, ethique animale, responsabilite et point de vue culturel.",
            "Lecture de documents historiques, proposition/argumentation, conditionnel, article/blog/proposal, revision Paper 1 et Paper 2.",
            "Paper 1: proposition/article/blog. Paper 2 Reading/Listening: mock-style checkpoints.",
        ),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Unite DP1", "Themes et contenus", "Langue et competences", "Evaluation / preuves"]
    for idx, h in enumerate(headers):
        add_cell_text(table.rows[0].cells[idx], h, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            add_cell_text(cells[idx], val)
    set_table_width(table, [2300, 2500, 2500, 2060])


def add_year2_plan(doc):
    add_heading(doc, "2. Plan DP2 2026-2027: aout-avril", 1)
    add_para(
        doc,
        "Le DP2 doit consolider les cinq themes IB tout en preparant explicitement les composantes d'examen. Les mocks doivent etre places juste apres les vacances de Noel; l'IA doit etre terminee en mars; avril doit rester consacre a la consolidation finale avant les examens DP de mai.",
    )

    rows = [
        (
            "Unite 1\nAout-mi-septembre",
            "Medias, technologies et creation",
            "Human ingenuity: communication and media, technology, entertainment/artistic expressions. Liens: identities (identite numerique).",
            "Analyse de medias francophones, fiabilite des sources, influence des reseaux, place de l'IA et de la creation.",
            "ATL: media literacy, research, critical thinking. Paper 2 Reading + Paper 1.",
            "Article, blog, critique/review, post reseau social, discours court.",
        ),
        (
            "Unite 2\nMi-septembre-novembre",
            "Societes, education et monde du travail",
            "Social organization: education, working world, social relationships, rules and order. Liens: identities (valeurs, ambitions).",
            "Debats sur l'ecole, l'emploi, l'engagement citoyen et les inegalites. Preparation a l'argumentation et au registre formel.",
            "ATL: communication, collaboration, self-management. Paper 1 + Oral IA practice.",
            "Courriel formel, lettre, rapport/compte rendu, proposition, interview.",
        ),
        (
            "Unite 3\nNovembre-decembre",
            "Environnement, droits et mondialisation",
            "Sharing the planet: environment, human rights, equality, globalization, ethics. Liens: social organization.",
            "Documents authentiques sur climat, justice sociale, urbanisation, consommation et responsabilite individuelle/collective.",
            "ATL: thinking, information literacy. Paper 2 Reading + Paper 2 Listening. Revision organisee avant mocks.",
            "Article, brochure/affiche, set of instructions, discours, blog argumentatif.",
        ),
        (
            "Mocks\nJuste apres Noel",
            "Point d'etape officiel",
            "Les cinq themes, selon les documents disponibles.",
            "Mock Paper 1, Paper 2 Reading et Paper 2 Listening; retour individuel et priorites IA.",
            "ATL: reflection, self-management.",
            "Relecture des types de textes et erreurs recurrentes.",
        ),
        (
            "Unite 4\nJanvier-fevrier",
            "Culture, memoire, traditions et voix francophones",
            "Experiences: customs and traditions, life stories. Human ingenuity: artistic expressions. Identities: beliefs and values.",
            "Memoire, patrimoine, arts, traditions, recits personnels et extraits litteraires ou culturels choisis.",
            "ATL: communication, cultural awareness. Oral IA + Paper 2 Reading.",
            "Discours, critique, article culturel, lettre personnelle, extrait litteraire accompagne.",
        ),
        (
            "Unite 5\nFevrier-avril",
            "IA finale et consolidation transversale",
            "Tous les themes: integration et transfert. Recommended topics repris selon les lacunes observees aux mocks.",
            "Banque de visuels, entrainements oraux chronometres, strategies Paper 1/Paper 2, consolidation lexicale et grammaticale.",
            "ATL: self-management, communication, reflection. IA en mars + Paper 1/Paper 2.",
            "Tous les types de textes majeurs: courriel, lettre, blog, article, discours, proposition, rapport, critique.",
        ),
    ]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Periode", "Unite", "Themes / topics", "Contenus", "Competences / examen", "Types de textes"]
    for idx, h in enumerate(headers):
        add_cell_text(table.rows[0].cells[idx], h, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            add_cell_text(cells[idx], val)
    set_table_width(table, [1100, 1550, 2150, 2150, 1550, 860])


def add_text_types_and_coverage(doc):
    add_heading(doc, "3. Couverture attendue", 1)
    add_heading(doc, "Types de textes Paper 1 a entretenir toute l'annee", 2)
    add_bullets(
        doc,
        [
            "Textes personnels: courriel, lettre personnelle, journal/blog, publication de reseaux sociaux.",
            "Textes professionnels ou formels: courriel formel, lettre, rapport/compte rendu, proposition, set of instructions.",
            "Textes medias: article, blog argumentatif, discours, critique/review, brochure/affiche, interview.",
            "Principe: les types de textes ne doivent pas etre bloques dans une seule unite. Chaque unite en pratique au moins un en production et en reconnait plusieurs en reception.",
        ],
    )

    add_heading(doc, "Themes IB et recommended topics encore a consolider", 2)
    rows = [
        ("Identities", "Bien entame: wellbeing, language and identity.", "Revenir par l'IA: valeurs, croyances, identite numerique, styles de vie."),
        ("Experiences", "Bien entame: life stories, rites of passage, migration.", "Approfondir customs/traditions, memoire, voyages, recits personnels."),
        ("Human ingenuity", "Introduit via arts, medias, IA oral.", "A traiter plus fortement: technology, communication and media, artistic expression, innovation."),
        ("Social organization", "Introduit par chomage et organisation sociale.", "Renforcer education, working world, social relationships, rules/order, civic life."),
        ("Sharing the planet", "Introduit par mondialisation, environnement urbain/rural, ethique.", "Renforcer environment, equality, human rights, globalization, ethical responsibility."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, h in enumerate(["Theme", "DP1 deja couvert", "DP2 a consolider"]):
        add_cell_text(table.rows[0].cells[idx], h, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            add_cell_text(cells[idx], val)
    set_table_width(table, [1700, 3800, 3860])

    add_heading(doc, "Point de vigilance", 2)
    add_para(
        doc,
        "La planification doit eviter de reproduire une progression du type theme 1, puis theme 2, puis theme 3. En Francais B SL, les eleves progressent mieux lorsque chaque unite combine contenu culturel, competence d'examen, type de texte, strategie ATL et retour metacognitif. C'est aussi la seule facon realiste de couvrir le programme sans sacrifier l'IA ni la preparation Paper 1/Paper 2.",
    )


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Francais B SL - Plan 2025-2027")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(89, 89, 89)


def main():
    doc = Document()
    style_document(doc)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Plan de cours sur deux ans - Francais B SL")
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    r.bold = True
    subtitle = doc.add_paragraph("DP1 2025-2026 et DP2 2026-2027 | Version de travail pour harmonisation IB")
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    add_info_table(doc)
    add_dp1_summary(doc)
    add_year2_plan(doc)
    add_text_types_and_coverage(doc)
    add_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
